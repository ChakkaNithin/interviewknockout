from contextlib import asynccontextmanager
from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, Form, Request
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pymongo.errors import PyMongoError
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
import os
import io
import json
import hmac
import hashlib
import logging
from pathlib import Path
from datetime import datetime, timezone
from typing import List, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env', override=True)

from models import (
    UserSignup, UserLogin, GoogleAuthRequest, UserPublic, AuthResponse,
    ResumeCreate, ResumeUpdate, Resume, ResumeData,
    ATSAnalyzeRequest, ATSResult,
    JDTailorRequest, JDTailorResult,
    AIGenerateRequest, AIGenerateResponse,
    JobSearchRequest, JobSearchResponse, Job,
    PaymentOrderRequest, PaymentOrderResponse, PaymentVerifyRequest,
    uid,
)
from auth import hash_password, verify_password, create_access_token, get_current_user_id
import ai_service
import jobs_service
from file_utils import extract_text_from_upload, extract_text_from_bytes

mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
client = AsyncIOMotorClient(mongo_url, serverSelectionTimeoutMS=5000)
db = client[os.environ.get('DB_NAME', 'interviewknockout')]

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

MAX_UPLOAD_BYTES = 5 * 1024 * 1024  # 5 MB
DEVELOPER_EMAILS = {"nithinnani1311@gmail.com"}

limiter = Limiter(key_func=get_remote_address)


@asynccontextmanager
async def lifespan(app: FastAPI):
    app.state.db_ready = False
    try:
        await db.users.create_index("email", unique=True)
        await db.resumes.create_index([("user_id", 1), ("updated_at", -1)])
        app.state.db_ready = True
        logger.info("Database connection ready")
    except Exception:
        logger.exception("Database initialization failed; API will stay up but database-backed routes may return 503")
    yield
    client.close()


app = FastAPI(title="InterviewKnockout API", version="1.0.0", lifespan=lifespan)

app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

app.add_middleware(
    CORSMiddleware,
    allow_origins=os.environ.get("CORS_ORIGINS", "http://localhost:3000").split(","),
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=["Authorization", "Content-Type", "Accept"],
)

api_router = APIRouter(prefix="/api")


def ensure_db_available() -> None:
    if not getattr(app.state, "db_ready", False):
        raise HTTPException(
            status_code=503,
            detail="Database is unavailable right now. Please try again in a moment.",
        )


def is_temporary_ai_error(error: Exception) -> bool:
    msg = str(error).lower()
    return any(
        marker in msg
        for marker in (
            "503",
            "unavailable",
            "high demand",
            "429",
            "resource_exhausted",
            "quota",
            "rate limit",
            "empty response",
        )
    )


def effective_plan_for_user(u: dict) -> str:
    email = (u.get("email") or "").lower()
    if email in DEVELOPER_EMAILS:
        return "premium"
    return u.get("plan", "free")


def user_to_public(u: dict) -> UserPublic:
    return UserPublic(
        id=u["id"],
        email=u["email"],
        name=u["name"],
        plan=effective_plan_for_user(u),
        avatar=u.get("avatar"),
        provider=u.get("provider", "email"),
        phone=u.get("phone"),
        location=u.get("location"),
        linkedin=u.get("linkedin"),
        github=u.get("github"),
        created_at=u.get("created_at", datetime.now(timezone.utc)),
    )


@api_router.get("/")
async def root():
    return {"name": "InterviewKnockout API", "status": "ok", "version": "1.0.0"}


@api_router.get("/health")
async def health():
    if not getattr(app.state, "db_ready", False):
        return {"status": "unhealthy", "error": "Database unreachable"}
    try:
        await db.command("ping")
        app.state.db_ready = True
        return {"status": "healthy", "db": "connected"}
    except Exception:
        app.state.db_ready = False
        logger.exception("Health check failed")
        return {"status": "unhealthy", "error": "Database unreachable"}


# ========== Auth ==========
@api_router.post("/auth/signup", response_model=AuthResponse)
@limiter.limit("5/minute")
async def signup(request: Request, payload: UserSignup):
    ensure_db_available()
    existing = await db.users.find_one({"email": payload.email.lower()})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    user_id = uid()
    doc = {
        "id": user_id,
        "email": payload.email.lower(),
        "name": payload.name,
        "password_hash": hash_password(payload.password),
        "plan": "free",
        "avatar": None,
        "provider": "email",
        "created_at": datetime.now(timezone.utc),
    }
    try:
        await db.users.insert_one(doc)
    except PyMongoError:
        raise HTTPException(status_code=400, detail="Email already registered")
    token = create_access_token(user_id)
    return AuthResponse(user=user_to_public(doc), token=token)


@api_router.post("/auth/login", response_model=AuthResponse)
@limiter.limit("10/minute")
async def login(request: Request, payload: UserLogin):
    ensure_db_available()
    u = await db.users.find_one({"email": payload.email.lower()})
    if not u or not verify_password(payload.password, u.get("password_hash", "")):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    token = create_access_token(u["id"])
    return AuthResponse(user=user_to_public(u), token=token)


@api_router.post("/auth/google", response_model=AuthResponse)
@limiter.limit("10/minute")
async def google_auth(request: Request, payload: GoogleAuthRequest):
    ensure_db_available()
    import asyncio
    from google.oauth2 import id_token as google_id_token
    from google.auth.transport import requests as google_requests

    google_client_id = os.environ.get('GOOGLE_CLIENT_ID')
    if not google_client_id:
        raise HTTPException(status_code=503, detail="Google auth not configured")

    try:
        # verify_oauth2_token is a blocking HTTP call — run it off the event loop
        loop = asyncio.get_event_loop()
        idinfo = await loop.run_in_executor(
            None,
            lambda: google_id_token.verify_oauth2_token(
                payload.token,
                google_requests.Request(),
                google_client_id,
            )
        )
    except ValueError as e:
        logger.warning(f"Google token verification failed: {e}")
        raise HTTPException(status_code=401, detail="Invalid Google token — please try signing in again")

    email = idinfo.get('email', '').lower()
    if not email:
        raise HTTPException(status_code=401, detail="Google account has no email address")

    if not idinfo.get('email_verified', False):
        raise HTTPException(status_code=401, detail="Google account email is not verified")

    name = idinfo.get('name', email.split('@')[0])
    google_id = idinfo['sub']
    picture = idinfo.get('picture')

    u = await db.users.find_one({"email": email})
    if not u:
        user_id = uid()
        doc = {
            "id": user_id,
            "email": email,
            "name": name,
            "password_hash": "",
            "plan": "free",
            "avatar": picture,
            "provider": "google",
            "google_id": google_id,
            "created_at": datetime.now(timezone.utc),
        }
        await db.users.insert_one(doc)
        u = doc
    else:
        # Keep avatar and google_id up to date on each login
        await db.users.update_one(
            {"email": email},
            {"$set": {"avatar": picture, "google_id": google_id}}
        )
        u = await db.users.find_one({"email": email})

    token = create_access_token(u["id"])
    return AuthResponse(user=user_to_public(u), token=token)


@api_router.get("/auth/me", response_model=UserPublic)
async def me(user_id: str = Depends(get_current_user_id)):
    ensure_db_available()
    u = await db.users.find_one({"id": user_id})
    if not u:
        raise HTTPException(status_code=404, detail="User not found")
    return user_to_public(u)


@api_router.patch("/auth/me/plan")
async def update_plan(plan: str, user_id: str = Depends(get_current_user_id)):
    ensure_db_available()
    if plan not in ("free",):
        # Paid plan upgrades require payment verification via Stripe webhook — not self-service
        raise HTTPException(status_code=402, detail="Plan upgrade requires payment")
    await db.users.update_one({"id": user_id}, {"$set": {"plan": plan}})
    u = await db.users.find_one({"id": user_id})
    return user_to_public(u)


# ========== Resumes ==========
@api_router.get("/resumes", response_model=List[Resume])
async def list_resumes(user_id: str = Depends(get_current_user_id)):
    ensure_db_available()
    cursor = db.resumes.find({"user_id": user_id}).sort("updated_at", -1)
    items = await cursor.to_list(200)
    result = []
    for r in items:
        r.pop("_id", None)
        result.append(Resume(**r))
    return result


@api_router.post("/resumes", response_model=Resume)
async def create_resume(payload: ResumeCreate, user_id: str = Depends(get_current_user_id)):
    ensure_db_available()
    r = Resume(
        user_id=user_id,
        title=payload.title,
        template=payload.template,
        target_role=payload.target_role,
        data=payload.data,
    )
    await db.resumes.insert_one(r.dict())
    return r


@api_router.get("/resumes/{resume_id}", response_model=Resume)
async def get_resume(resume_id: str, user_id: str = Depends(get_current_user_id)):
    ensure_db_available()
    r = await db.resumes.find_one({"id": resume_id, "user_id": user_id})
    if not r:
        raise HTTPException(status_code=404, detail="Resume not found")
    r.pop("_id", None)
    return Resume(**r)


@api_router.put("/resumes/{resume_id}", response_model=Resume)
async def update_resume(resume_id: str, payload: ResumeUpdate, user_id: str = Depends(get_current_user_id)):
    ensure_db_available()
    update = {k: v for k, v in payload.dict(exclude_none=True).items()}
    update["updated_at"] = datetime.now(timezone.utc)
    result = await db.resumes.update_one({"id": resume_id, "user_id": user_id}, {"$set": update})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Resume not found")
    r = await db.resumes.find_one({"id": resume_id})
    r.pop("_id", None)
    return Resume(**r)


@api_router.delete("/resumes/{resume_id}")
async def delete_resume(resume_id: str, user_id: str = Depends(get_current_user_id)):
    ensure_db_available()
    result = await db.resumes.delete_one({"id": resume_id, "user_id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Resume not found")
    return {"deleted": True}


# ========== ATS Analysis ==========
@api_router.post("/ats/analyze")
@limiter.limit("20/minute")
async def ats_analyze_text(request: Request, payload: ATSAnalyzeRequest, user_id: str = Depends(get_current_user_id)):
    try:
        result = await ai_service.analyze_ats(payload.resume_text, payload.target_role or "")
        return result
    except Exception as e:
        logger.exception("ATS analysis failed")
        if is_temporary_ai_error(e):
            raise HTTPException(status_code=503, detail="AI service is temporarily unavailable. Please try again in a moment.")
        raise HTTPException(status_code=500, detail="Analysis failed. Please try again.")


@api_router.post("/ats/analyze-file")
@limiter.limit("20/minute")
async def ats_analyze_file(
    request: Request,
    file: UploadFile = File(...),
    target_role: Optional[str] = Form(None),
    user_id: str = Depends(get_current_user_id),
):
    if file.size and file.size > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (max 5MB)")
    content = await file.read(MAX_UPLOAD_BYTES + 1)
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (max 5MB)")
    text = extract_text_from_bytes(content, file.filename or "")
    if len(text.strip()) < 100:
        raise HTTPException(status_code=400, detail="Could not extract sufficient text from resume")
    try:
        result = await ai_service.analyze_ats(text, target_role or "")
        result["filename"] = file.filename
        result["extracted_chars"] = len(text)
        result["original_resume_text"] = text
        return result
    except Exception as e:
        logger.exception("ATS analysis failed")
        if is_temporary_ai_error(e):
            raise HTTPException(status_code=503, detail="AI service is temporarily unavailable. Please try again in a moment.")
        raise HTTPException(status_code=500, detail="Analysis failed. Please try again.")


# ========== ATS Apply Fixes ==========
@api_router.post("/ats/apply-fixes")
@limiter.limit("10/minute")
async def ats_apply_fixes(request: Request, payload: dict, user_id: str = Depends(get_current_user_id)):
    resume_text = payload.get("resume_text", "")
    fixes = payload.get("fixes", [])
    target_role = payload.get("target_role", "")
    if len(resume_text.strip()) < 100:
        raise HTTPException(status_code=400, detail="Resume text is too short.")
    if not fixes:
        raise HTTPException(status_code=400, detail="No fixes provided.")
    try:
        result = await ai_service.apply_ats_fixes(resume_text, fixes, target_role)
        return result
    except Exception as e:
        logger.exception("Apply ATS fixes failed")
        if is_temporary_ai_error(e):
            raise HTTPException(status_code=503, detail="AI service is temporarily unavailable. Please try again in a moment.")
        raise HTTPException(status_code=500, detail="Failed to apply fixes. Please try again.")


# ========== JD Tailor ==========
@api_router.post("/jd/tailor")
@limiter.limit("20/minute")
async def jd_tailor(request: Request, payload: JDTailorRequest, user_id: str = Depends(get_current_user_id)):
    if len(payload.jd_text.strip()) < 50:
        raise HTTPException(status_code=400, detail="JD must be at least 50 characters")
    if len(payload.resume_text.strip()) < 100:
        raise HTTPException(status_code=400, detail="Resume text must be at least 100 characters")
    try:
        return await ai_service.tailor_to_jd(payload.resume_text, payload.jd_text)
    except Exception as e:
        logger.exception("JD tailoring failed")
        if is_temporary_ai_error(e):
            raise HTTPException(status_code=503, detail="AI service is temporarily unavailable. Please try again in a moment.")
        raise HTTPException(status_code=500, detail="Analysis failed. Please try again.")


@api_router.post("/jd/tailor-file")
@limiter.limit("20/minute")
async def jd_tailor_file(
    request: Request,
    jd_text: str = Form(...),
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user_id),
):
    if file.size and file.size > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (max 5MB)")
    content = await file.read(MAX_UPLOAD_BYTES + 1)
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (max 5MB)")
    text = extract_text_from_bytes(content, file.filename or "")
    if len(jd_text.strip()) < 50:
        raise HTTPException(status_code=400, detail="JD must be at least 50 characters")
    try:
        result = await ai_service.tailor_to_jd(text, jd_text)
        result["original_resume_text"] = text
        return result
    except Exception as e:
        logger.exception("JD tailoring failed")
        if is_temporary_ai_error(e):
            raise HTTPException(status_code=503, detail="AI service is temporarily unavailable. Please try again in a moment.")
        raise HTTPException(status_code=500, detail="Analysis failed. Please try again.")


@api_router.post("/jd/download-docx")
@limiter.limit("20/minute")
async def jd_download_docx(request: Request, payload: dict, user_id: str = Depends(get_current_user_id)):
    name          = payload.get("candidate_name", "Your Name")
    job_title     = payload.get("job_title", "")
    company       = payload.get("company", "")
    summary       = payload.get("tailored_summary", "")
    skills        = payload.get("tailored_skills", [])
    original_text = payload.get("original_resume_text", "")

    doc = Document()

    # Remove default margins for tighter layout
    for section in doc.sections:
        section.top_margin    = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin   = Pt(54)
        section.right_margin  = Pt(54)

    # Name
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x0F, 0x3D, 0x2E)
    p.paragraph_format.space_after = Pt(2)

    # Job title line
    if job_title:
        p2 = doc.add_paragraph()
        run2 = p2.add_run(f"{job_title}" + (f"  ·  {company}" if company else ""))
        run2.bold = True
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(0x0F, 0x3D, 0x2E)
        p2.paragraph_format.space_after = Pt(4)

    # Separator
    sep = doc.add_paragraph()
    sep.add_run("─" * 72).font.color.rgb = RGBColor(0x0F, 0x3D, 0x2E)
    sep.paragraph_format.space_after = Pt(6)

    # Tailored Professional Summary
    h1 = doc.add_paragraph()
    r1 = h1.add_run("PROFESSIONAL SUMMARY")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x0F, 0x3D, 0x2E)
    h1.paragraph_format.space_after = Pt(2)
    doc.add_paragraph(summary).paragraph_format.space_after = Pt(8)

    # Tailored Key Skills
    h2 = doc.add_paragraph()
    r2 = h2.add_run("KEY SKILLS")
    r2.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x0F, 0x3D, 0x2E)
    h2.paragraph_format.space_after = Pt(2)
    doc.add_paragraph(", ".join(skills)).paragraph_format.space_after = Pt(8)

    # Remaining sections from original resume (skip Summary + Skills)
    skip_titles = {"PROFESSIONAL SUMMARY", "SUMMARY", "OBJECTIVE", "CAREER OBJECTIVE",
                   "KEY SKILLS", "SKILLS", "TECHNICAL SKILLS", "CORE COMPETENCIES", "CORE SKILLS"}
    lines = original_text.split("\n")
    in_skip = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        upper = stripped.upper()
        if upper in skip_titles:
            in_skip = True
            continue
        # Detect new section header: short ALL-CAPS line
        if stripped == stripped.upper() and 3 <= len(stripped) <= 60 and stripped.replace(" ", "").isalpha():
            in_skip = False
            hx = doc.add_paragraph()
            rx = hx.add_run(stripped)
            rx.bold = True
            rx.font.size = Pt(10)
            rx.font.color.rgb = RGBColor(0x0F, 0x3D, 0x2E)
            hx.paragraph_format.space_before = Pt(8)
            hx.paragraph_format.space_after  = Pt(2)
            continue
        if not in_skip:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_title = (job_title or "tailored-resume").lower().replace(" ", "-").replace("/", "-")[:40]
    filename = f"{safe_title}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ========== AI Content Generation ==========
@api_router.post("/ai/generate", response_model=AIGenerateResponse)
@limiter.limit("20/minute")
async def ai_generate(request: Request, payload: AIGenerateRequest, user_id: str = Depends(get_current_user_id)):
    try:
        result = await ai_service.generate_content(payload.prompt, payload.context or "", payload.kind)
        return AIGenerateResponse(text=result.get("text", ""), suggestions=result.get("suggestions", []))
    except Exception:
        logger.exception("AI generate failed")
        raise HTTPException(status_code=500, detail="AI generation failed. Please try again.")


# ========== Job Search ==========
@api_router.post("/jobs/search", response_model=JobSearchResponse)
@limiter.limit("15/minute")
async def search_jobs_endpoint(request: Request, payload: JobSearchRequest, user_id: str = Depends(get_current_user_id)):
    try:
        raw = await jobs_service.search_jobs(
            search_term=payload.search_term,
            location=payload.location,
            hours_old=payload.hours_old,
            is_remote=payload.is_remote,
            results_per_page=payload.results_per_page,
            resume_skills=payload.resume_skills,
            min_salary=payload.min_salary,
            max_salary=payload.max_salary,
            seniority_levels=payload.seniority_levels,
            technologies=payload.technologies,
            min_employees=payload.min_employees,
            max_employees=payload.max_employees,
            min_funding=payload.min_funding,
            max_funding=payload.max_funding,
            industries=payload.industries,
            company_names=payload.company_names,
            easy_apply=payload.easy_apply,
            page=payload.page,
        )
        jobs: List[Job] = []
        for j in raw:
            try:
                jobs.append(Job(**j))
            except Exception:
                continue
        return JobSearchResponse(jobs=jobs, total=len(jobs), search_term=payload.search_term, location=payload.location)
    except Exception:
        logger.exception("Job search failed")
        raise HTTPException(status_code=500, detail="Job search failed. Please try again.")


# ========== Saved Jobs ==========
@api_router.post("/jobs/save")
async def save_job(job: Job, user_id: str = Depends(get_current_user_id)):
    doc = job.dict()
    doc["user_id"] = user_id
    doc["saved_at"] = datetime.now(timezone.utc)
    await db.saved_jobs.insert_one(doc)
    return {"saved": True, "id": job.id}


@api_router.get("/jobs/saved", response_model=List[Job])
async def get_saved_jobs(user_id: str = Depends(get_current_user_id)):
    cursor = db.saved_jobs.find({"user_id": user_id}).sort("saved_at", -1)
    items = await cursor.to_list(200)
    result = []
    for j in items:
        j.pop("_id", None)
        j.pop("user_id", None)
        j.pop("saved_at", None)
        try:
            result.append(Job(**j))
        except Exception:
            continue
    return result


@api_router.delete("/jobs/saved/{job_id}")
async def unsave_job(job_id: str, user_id: str = Depends(get_current_user_id)):
    result = await db.saved_jobs.delete_one({"user_id": user_id, "id": job_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Saved job not found")
    return {"deleted": True}


# ========== Payments (Razorpay) ==========

# Prices in paise (INR). 1 INR = 100 paise.
# Update these when the owner confirms final pricing.
PLAN_PRICES = {
    "pro_monthly":      99900,    # ₹999/month
    "pro_yearly":      719900,    # ₹7199/year  (~₹600/month, 40% off)
    "premium_monthly": 199900,    # ₹1999/month
    "premium_yearly":  143900 * 12 // 10,  # ₹17268/year (~₹1439/month)
}


@api_router.post("/payments/create-order", response_model=PaymentOrderResponse)
@limiter.limit("5/minute")
async def create_payment_order(request: Request, payload: PaymentOrderRequest, user_id: str = Depends(get_current_user_id)):
    key_id = os.environ.get("RAZORPAY_KEY_ID", "")
    key_secret = os.environ.get("RAZORPAY_KEY_SECRET", "")

    if not key_id or key_id.startswith("REPLACE"):
        raise HTTPException(status_code=503, detail="Payment gateway not configured")

    if payload.plan not in ("pro", "premium"):
        raise HTTPException(status_code=400, detail="Invalid plan")
    if payload.billing not in ("monthly", "yearly"):
        raise HTTPException(status_code=400, detail="Invalid billing period")

    amount = PLAN_PRICES.get(f"{payload.plan}_{payload.billing}", 0)

    try:
        import razorpay
        rzp = razorpay.Client(auth=(key_id, key_secret))
        order = rzp.order.create({
            "amount": amount,
            "currency": "INR",
            "receipt": f"{user_id[:8]}-{payload.plan}",
            "notes": {
                "plan": payload.plan,
                "billing": payload.billing,
                "user_id": user_id,
            },
        })
        return PaymentOrderResponse(
            order_id=order["id"],
            amount=amount,
            currency="INR",
            key_id=key_id,
        )
    except Exception:
        logger.exception("Razorpay order creation failed")
        raise HTTPException(status_code=500, detail="Could not initiate payment. Please try again.")


@api_router.post("/payments/verify")
async def verify_payment(payload: PaymentVerifyRequest, user_id: str = Depends(get_current_user_id)):
    key_secret = os.environ.get("RAZORPAY_KEY_SECRET", "")

    # Verify HMAC-SHA256 signature
    msg = f"{payload.razorpay_order_id}|{payload.razorpay_payment_id}"
    expected = hmac.new(key_secret.encode(), msg.encode(), hashlib.sha256).hexdigest()
    if not hmac.compare_digest(expected, payload.razorpay_signature):
        raise HTTPException(status_code=400, detail="Payment verification failed — invalid signature")

    if payload.plan not in ("pro", "premium"):
        raise HTTPException(status_code=400, detail="Invalid plan")

    # Upgrade plan
    await db.users.update_one({"id": user_id}, {"$set": {"plan": payload.plan}})

    # Record payment
    await db.payments.insert_one({
        "user_id": user_id,
        "order_id": payload.razorpay_order_id,
        "payment_id": payload.razorpay_payment_id,
        "plan": payload.plan,
        "billing": payload.billing,
        "amount": PLAN_PRICES.get(f"{payload.plan}_{payload.billing}", 0),
        "currency": "INR",
        "created_at": datetime.now(timezone.utc),
    })

    u = await db.users.find_one({"id": user_id})
    return user_to_public(u)


@api_router.post("/payments/webhook")
async def razorpay_webhook(request: Request):
    body = await request.body()
    webhook_secret = os.environ.get("RAZORPAY_WEBHOOK_SECRET", "")
    sig = request.headers.get("X-Razorpay-Signature", "")

    if not webhook_secret or webhook_secret.startswith("REPLACE"):
        raise HTTPException(status_code=503, detail="Webhook not configured")

    expected = hmac.new(webhook_secret.encode(), body, hashlib.sha256).hexdigest()
    if not hmac.compare_digest(expected, sig):
        raise HTTPException(status_code=400, detail="Invalid webhook signature")

    try:
        event = json.loads(body)
        event_type = event.get("event", "")
        logger.info(f"Razorpay webhook received: {event_type}")

        if event_type == "payment.captured":
            notes = event.get("payload", {}).get("payment", {}).get("entity", {}).get("notes", {})
            uid_val = notes.get("user_id")
            plan = notes.get("plan")
            if uid_val and plan in ("pro", "premium"):
                await db.users.update_one({"id": uid_val}, {"$set": {"plan": plan}})
                logger.info(f"Webhook: upgraded user {uid_val} to {plan}")
    except Exception:
        logger.exception("Webhook processing error")

    return {"status": "ok"}


app.include_router(api_router)
