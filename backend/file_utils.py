import io
import logging
from fastapi import UploadFile, HTTPException

logger = logging.getLogger(__name__)


async def extract_text_from_upload(file: UploadFile) -> str:
    content = await file.read()
    return extract_text_from_bytes(content, file.filename or "")


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    name = filename.lower()
    if not name.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported. Please upload a Word document.")
    return _extract_docx(content)


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts).strip()
    except Exception as e:
        logger.exception("DOCX extraction failed")
        raise HTTPException(status_code=400, detail=f"Failed to read DOCX: {e}")
